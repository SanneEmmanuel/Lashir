import os
import io
import tempfile
import requests
import librosa
import numpy as np
from flask import Flask, request, send_file, jsonify
import music21
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib import patches
import matplotlib.lines as mlines
from flask_cors import CORS

# Sanne Flask Initialization
app = Flask(__name__)

# Enable CORS for all routes and all origins, Edit Later
CORS(app)
def root_redirect():
    return redirect("https://github.com/sanneemmanuel/libra", code=302)

app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configuration
FLAT_IO_API_KEY = os.getenv('FLAT_IO_API_KEY', '')
APP_TITLE = "Lashir: Your Ultimate Sound to Solfa Converter"
AUTHOR = "Dr. Sanne Karibo"

# Chromatic solfege mapping
SOLFEGE_MAP = {
    0: 'Do', 1: 'Di', 2: 'Re', 3: 'Ri', 4: 'Mi',
    5: 'Fa', 6: 'Fi', 7: 'Sol', 8: 'Si', 9: 'La',
    10: 'Li', 11: 'Ti'
}

# Traditional sheet music styling
STAFF_COLOR = '#000000'
NOTE_COLOR = '#000000'
BACKGROUND_COLOR = '#FFFFFF'
STAFF_LINE_WIDTH = 1.2
NOTE_RADIUS = 0.4
STEM_WIDTH = 0.1
STEM_LENGTH = 3.5

def setup_sheet_music_style():
    """Configure traditional sheet music styling"""
    plt.rcParams.update({
        'figure.facecolor': BACKGROUND_COLOR,
        'axes.facecolor': BACKGROUND_COLOR,
        'axes.edgecolor': 'none',
        'axes.linewidth': 0,
        'grid.color': 'none',
        'font.family': 'serif',
        'font.serif': ['Times New Roman', 'Times', 'serif'],
        'font.size': 12,
        'savefig.dpi': 300,
        'savefig.bbox': 'tight',
        'savefig.pad_inches': 0.1
    })

def audio_to_notes(audio_path):
    """Convert audio to professional music notation"""
    y, sr = librosa.load(audio_path, sr=None)
    
    # Professional pitch detection
    tempo, _ = librosa.beat.beat_track(y=y, sr=sr)
    pitches, magnitudes = librosa.piptrack(y=y, sr=sr, fmin=50, fmax=2000)
    onset_frames = librosa.onset.onset_detect(
        y=y, sr=sr, units='frames',
        pre_max=3, post_max=3,
        pre_avg=3, post_avg=5,
        delta=0.2, wait=10
    )
    onset_times = librosa.frames_to_time(onset_frames, sr=sr)
    
    notes = []
    for start_time in onset_times:
        frame_idx = np.argmin(np.abs(librosa.times_like(pitches) - start_time))
        pitch_idx = np.argmax(magnitudes[:, frame_idx])
        pitch_hz = pitches[pitch_idx, frame_idx]
        
        if pitch_hz > 0:
            midi = librosa.hz_to_midi(pitch_hz)
            pitch = music21.pitch.Pitch(midi=midi)
            semitone = int(round(midi)) % 12
            solfege = SOLFEGE_MAP.get(semitone, '?')
            
            note = music21.note.Note(pitch)
            note.duration = music21.duration.Duration(0.5)
            note.addLyric(solfege)
            notes.append(note)
    
    return notes, tempo

def create_staff_notation(notes, tempo):
    """Create professional music staff"""
    stream = music21.stream.Stream()
    stream.append(music21.tempo.MetronomeMark(number=tempo))
    stream.append(music21.meter.TimeSignature('4/4'))
    
    for note in notes:
        stream.append(note)
    
    return stream

def render_traditional_staff(stream, img_path):
    """Render professional sheet music with Matplotlib"""
    try:
        setup_sheet_music_style()
        fig, ax = plt.subplots(figsize=(12, 4))
        
        # Draw staff lines
        for i in range(-2, 3):
            ax.axhline(y=i, color=STAFF_COLOR, linewidth=STAFF_LINE_WIDTH, zorder=1)
        
        # Draw clef
        treble_clef = r"$\trebleclef$"
        ax.text(0.2, 0.5, treble_clef, fontsize=48, 
                ha='center', va='center', transform=ax.transAxes)
        
        # Draw notes
        x_pos = 0.3
        note_spacing = 0.15
        
        for note in stream.flat.notes:
            # Calculate staff position (middle C = 0, each line/space = 0.5)
            staff_pos = (note.pitch.midi - 60) / 2
            
            # Draw note head
            note_head = patches.Circle(
                (x_pos, staff_pos), 
                radius=NOTE_RADIUS,
                facecolor=NOTE_COLOR,
                edgecolor=NOTE_COLOR,
                zorder=2
            )
            ax.add_patch(note_head)
            
            # Draw stem
            stem_direction = -1 if staff_pos >= 0 else 1
            stem = mlines.Line2D(
                [x_pos + NOTE_RADIUS, x_pos + NOTE_RADIUS],
                [staff_pos, staff_pos + stem_direction * STEM_LENGTH],
                color=NOTE_COLOR,
                linewidth=STEM_WIDTH,
                zorder=2
            )
            ax.add_line(stem)
            
            # Add solfege label
            ax.text(
                x_pos, staff_pos - (1.5 * stem_direction),
                note.lyric,
                ha='center', va='center',
                fontsize=10, fontweight='bold'
            )
            
            x_pos += note_spacing
        
        # Set professional boundaries
        ax.set_xlim(0, x_pos + 0.5)
        ax.set_ylim(-4.5, 4.5)
        ax.set_aspect('equal')
        ax.axis('off')
        
        # Add metadata
        fig.text(0.5, 0.95, APP_TITLE, 
                ha='center', va='top', fontsize=14, fontweight='bold')
        fig.text(0.5, 0.90, f"Tempo: {tempo} BPM | Key: C Major", 
                ha='center', va='top', fontsize=10)
        fig.text(0.5, 0.05, f"Generated by {AUTHOR}", 
                ha='center', va='bottom', fontsize=9)
        
        plt.savefig(img_path)
        plt.close(fig)
        return True
    except Exception as e:
        print(f"Professional rendering failed: {str(e)}")
        return False

def render_with_webapi(stream, img_path):
    """Render via Flat.io API"""
    if not FLAT_IO_API_KEY:
        return False
        
    try:
        # Convert to MusicXML
        xml = stream.write('musicxml', fp=io.BytesIO()).getvalue()
        
        # Create score
        create_res = requests.post(
            'https://api.flat.io/v2/scores',
            headers={'Authorization': f'Bearer {FLAT_IO_API_KEY}'},
            files={'file': ('notation.xml', xml)},
            data={'title': 'Lashir Notation', 'privacy': 'private'}
        )
        score_id = create_res.json()['id']
        
        # Get PNG image
        img_res = requests.get(
            f'https://api.flat.io/v2/scores/{score_id}/png',
            headers={'Authorization': f'Bearer {FLAT_IO_API_KEY}'},
            params={'resolution': 300, 'margin': 50}
        )
        
        with open(img_path, 'wb') as f:
            f.write(img_res.content)
        
        # Cleanup
        requests.delete(
            f'https://api.flat.io/v2/scores/{score_id}',
            headers={'Authorization': f'Bearer {FLAT_IO_API_KEY}'}
        )
            
        return True
    except Exception as e:
        print(f"Web API error: {str(e)}")
        return False

def create_professional_doc(notes, tempo, img_path=None):
    """Create Word document with professional styling"""
    doc = Document()
    
    # Title page
    title_section = doc.sections[0]
    title_section.different_first_page_header_footer = True
    header = title_section.first_page_header
    header.is_linked_to_previous = False
    
    # Custom title page design
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(APP_TITLE)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Music Transcription Report", style='Heading1').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    author_para = doc.add_paragraph()
    author_run = author_para.add_run(f"Generated by {AUTHOR}")
    author_run.font.size = Pt(14)
    author_run.font.italic = True
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Tempo: {tempo} BPM | Key: C Major")
    doc.add_page_break()
    
    # Staff notation section
    if img_path and os.path.exists(img_path):
        doc.add_heading('Staff Notation', level=1)
        doc.add_picture(img_path, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 1: Professional staff notation with solfege syllables")
    else:
        doc.add_heading('Staff Notation Unavailable', level=1)
        doc.add_paragraph("Professional notation rendering failed. Below is your solfege analysis:")
    
    # Solfege guide with professional table
    doc.add_heading('Solfa Analysis', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'LightShading-Accent1'
    
    # Table header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Note'
    hdr_cells[1].text = 'Solfa Syllable'
    hdr_cells[2].text = 'Frequency (Hz)'
    
    # Group notes by pitch
    pitch_groups = {}
    for note in notes:
        pitch_name = note.pitch.name
        if pitch_name not in pitch_groups:
            pitch_groups[pitch_name] = {
                'solfa': note.lyric,
                'freq': note.pitch.frequency,
                'count': 0
            }
        pitch_groups[pitch_name]['count'] += 1
    
    # Add table rows sorted by frequency
    for pitch_name, data in sorted(pitch_groups.items(), 
                                 key=lambda x: x[1]['freq']):
        row_cells = table.add_row().cells
        row_cells[0].text = pitch_name
        row_cells[1].text = data['solfa']
        row_cells[2].text = f"{data['freq']:.1f} Hz"
    
    # Add analysis summary
    doc.add_heading('Performance Analysis', level=2)
    doc.add_paragraph(f"Total notes detected: {len(notes)}")
    doc.add_paragraph(f"Pitch range: {min(pitch_groups.keys())} to {max(pitch_groups.keys())}")
    
    # Add footer with copyright
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.text = f"© {AUTHOR} | {APP_TITLE} | Confidential Report"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def save_document(doc, output_path):
    """Save document with error handling"""
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Document save failed: {str(e)}")
        return False

def process_audio(audio_path):
    """Full audio processing pipeline"""
    try:
        # Process audio
        notes, tempo = audio_to_notes(audio_path)
        stream = create_staff_notation(notes, tempo)
        
        # Create temp image path
        img_path = os.path.join(tempfile.gettempdir(), 'staff_notation.png')
        
        # Try rendering methods
        if not render_traditional_staff(stream, img_path):
            if not render_with_webapi(stream, img_path):
                img_path = None
        
        # Create professional document
        doc = create_professional_doc(notes, tempo, img_path)
        return doc
    except Exception as e:
        print(f"Processing error: {str(e)}")
        return None

@app.route('/upload', methods=['POST'])
def upload():
    """Handle file upload"""
    if 'file' not in request.files:
        return jsonify(error="No file part"), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify(error="No selected file"), 400
    
    try:
        # Save audio
        audio_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(audio_path)
        
        # Process audio
        doc = process_audio(audio_path)
        if not doc:
            return jsonify(error="Audio processing failed"), 500
        
        # Save document
        doc_path = os.path.join(app.config['UPLOAD_FOLDER'], 'Lashir_Notation.docx')
        if not save_document(doc, doc_path):
            return jsonify(error="Document creation failed"), 500
        
        return send_file(
            doc_path,
            as_attachment=True,
            download_name='Lashir_Professional_Notation.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify(error=f"Processing error: {str(e)}"), 500

@app.route('/stream', methods=['POST'])
def stream():
    """Handle audio stream"""
    audio_data = request.data
    if not audio_data:
        return jsonify(error="No audio data received"), 400
    
    try:
        # Save to temp file
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp:
            tmp.write(audio_data)
            audio_path = tmp.name
        
        # Process audio
        doc = process_audio(audio_path)
        if not doc:
            return jsonify(error="Audio processing failed"), 500
        
        # Save document
        doc_path = os.path.join(tempfile.gettempdir(), 'Lashir_Stream_Notation.docx')
        if not save_document(doc, doc_path):
            return jsonify(error="Document creation failed"), 500
        
        return send_file(
            doc_path,
            as_attachment=True,
            download_name='Lashir_Professional_Stream_Notation.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify(error=f"Processing error: {str(e)}"), 500

@app.route('/')
def index():
    return jsonify(
        application=APP_TITLE,
        author=AUTHOR,
        endpoints=["/upload", "/stream"],
        status="Ready"
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
